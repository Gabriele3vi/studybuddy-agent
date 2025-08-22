'''
Questo script deve esporre le funzioni per inserire un nuovo documento all'interno di TiDB.

Requisiti:
- Riconoscimetnto tipo di documento
- conversione in testo
- pulizia del testo
- creazione chunk per l'embedding
- crezione dell'embedding
- caricamento dei chunck all'interno del database
'''
import PyPDF2
from docx import Document
from openai import OpenAI
import os 
from dotenv import load_dotenv
import tiktoken
import pymysql
import pymysql.cursors
from pathlib import Path
import datetime as dt
import numpy as np
import ssl

ctx = ssl.create_default_context()


load_dotenv()

def get_db_connection():
    connection = pymysql.connect(
        host = os.getenv('tidb_host'),
        port = int(os.getenv('tidb_port')),
        user = os.getenv('tidb_user'),
        password = os.getenv('tidb_pass'),
        database = os.getenv('tidb_db'),
        ssl_verify_cert = True,
        ssl_verify_identity = True
    )

    return connection

def extract_text_from_pdf(filepath):
    # Open the PDF file
    text = ""
    with open(filepath, "rb") as file:
        reader = PyPDF2.PdfReader(file)

        for page in reader.pages:
            text += page.extract_text() + " "
        
    return text


def extract_text_from_docx(filepath):
    doc = Document(filepath)
    text = ""

    # Extract from paragraphs
    for para in doc.paragraphs:
        text += para.text + " "

    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            text += "\t".join(row_text) + "\n"

    return text



def chunk_text(text, openai_client, max_tokens=400, overlap=50, model="text-embedding-3-small"):
    # Keywords è un vettore di parole chiave che mettiamo sempre all'interno di un testo per aiutare la vector similarity
    # Chunko i risultati
    enc = tiktoken.encoding_for_model(model)
    tokens = enc.encode(text)

    chunks = []
    start = 0
    while start < len(tokens):
        end = start + max_tokens
        chunk = enc.decode(tokens[start:end])
        chunks.append(chunk)
        start += max_tokens - overlap


    # Faccio gli embeddings
    embedded_chunks = []
    for chk in chunks:
        response = openai_client.embeddings.create(
            input=chk,
            model="text-embedding-3-small"
        )
        embedded_chunks.append((chk, str(response.data[0].embedding)))

    return embedded_chunks



def upload_file(connection, openai_client, filepath):
    filename, ext = os.path.splitext(os.path.basename(filepath))
    if ext == '.pdf':
        text = extract_text_from_pdf(filepath)
    elif ext == '.docx':
        text = extract_text_from_docx(filepath)
    embedded_chunks = chunk_text(text, openai_client)
    
    try:
        with connection:
            with connection.cursor() as cursor:
                # Create a new record
                upload_file_sql = "INSERT INTO `notes` (`name`, `ext`) VALUES (%s, %s)"
                cursor.execute(upload_file_sql, (filename, ext))
                note_id = cursor.lastrowid

                # Insert into notes
                chk_num = 0
                for content, embedding in embedded_chunks:
                    chk_num += 1
                    sql = "INSERT INTO `note_chunks` (`note_id`, `chunk_order`, `content`, `embedding`) VALUES (%s, %s, %s, %s)"
                    cursor.execute(sql, (note_id, chk_num, content, embedding))
                    connection.commit()

            connection.commit()
            return "File caricato correttamente"
    except Exception as e:
        connection.rollback()
        print("Transaction rolled back:", e)
        print("Non sono riuscto a caricare il file")

    



if __name__ == '__main__':
    load_dotenv()

    OPENAI_API_KEY = os.getenv('openai_api_key')
    PDF_TEST_FILE = 'docs/Postgre commands.pdf'
    DOCX_TEST_FILE = 'docs/Passuello fratelli intro.docx'

    connection = get_db_connection()

    openai_client = OpenAI(api_key=OPENAI_API_KEY)
    

    filepath="docs/Postgre commands.pdf"
    subject = 'sql'

    file_id = upload_file(connection, openai_client, filepath, subject)
    print(f"File added with ID {file_id}")
    
    

    #text = extract_text_from_docx(DOCX_TEST_FILE)
    #chks = chunk_text(text)

    #ingest_chunks(client, None, 1, chks)